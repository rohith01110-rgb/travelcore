from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

doc = Document()

# Add Title
title = doc.add_heading('Infera Core - Technical Architecture & Scalability Report', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add Logo (if exists)
logo_path = r'd:\projects\inferacore\frontend\public\logo.png'
if os.path.exists(logo_path):
    doc.add_picture(logo_path, width=Inches(3.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph('This report provides a deep technical analysis of the Infera Core platform, an intelligent engineering workspace and Neural Study engine. The platform utilizes a decoupled, modern architecture separating a Next.js (React) frontend from a FastAPI (Python) AI-agent backend, orchestrated with Supabase for data management and authentication.')

doc.add_heading('2. System Architecture', level=1)
doc.add_paragraph('The architecture follows a decoupled client-server model, optimized for heavy AI processing, advanced Retrieval-Augmented Generation (RAG), and real-time user interactions.')

doc.add_heading('2.1 Frontend Architecture', level=2)
doc.add_paragraph('The frontend is built for extreme performance, rich user experience, and interactive AI widgets.')
doc.add_paragraph('- Framework: Next.js 16 (App Router) & React 19', style='List Bullet')
doc.add_paragraph('- Language: TypeScript for strict type-safety', style='List Bullet')
doc.add_paragraph('- Styling: Tailwind CSS 4, Framer Motion, Radix UI, Sonner', style='List Bullet')
doc.add_paragraph('- AI Integration: Vercel AI SDK (@ai-sdk/google, @ai-sdk/groq)', style='List Bullet')

doc.add_heading('2.2 Backend Architecture', level=2)
doc.add_paragraph('The backend serves as the heavy-lifting engine for AI agents and complex orchestration.')
doc.add_paragraph('- Framework: FastAPI running on Python 3.13', style='List Bullet')
doc.add_paragraph('- Orchestration: LangChain and LangGraph', style='List Bullet')
doc.add_paragraph('- LLMs: Google GenAI, OpenAI, and Groq SDKs', style='List Bullet')
doc.add_paragraph('- Search/RAG: Tavily, mixedbread, SQLAlchemy, vecs (pgvector)', style='List Bullet')
doc.add_paragraph('- Multi-modal: Vision endpoints processing Markdown and uploads', style='List Bullet')

doc.add_heading('3. Scalability Analysis', level=1)
doc.add_paragraph('The platform is designed to be highly scalable, addressing AI application bottlenecks.')
doc.add_paragraph('- Horizontal Scalability: Stateless FastAPI backend and Edge-ready Next.js frontend.', style='List Bullet')
doc.add_paragraph('- Database: Supabase connection pooling and direct vector storage with vecs.', style='List Bullet')
doc.add_paragraph('- Latency: Streaming responses via Vercel AI SDK and Multi-Model fallbacks.', style='List Bullet')

doc.add_heading('4. Core Features Deep-Dive', level=1)
doc.add_heading('4.1 Advanced AI Orchestration', level=2)
doc.add_paragraph('- General Agent (/chat): Multi-purpose coding & assistance.', style='List Bullet')
doc.add_paragraph('- Study Agent (/study): Academic use with Deep Think and Web Search.', style='List Bullet')
doc.add_paragraph('- Resume Agent (/resume-analyze): Specialized pipeline parsing PDFs with PyPDF2.', style='List Bullet')

doc.add_heading('4.2 Multi-Modal & Dashboard Integration', level=2)
doc.add_paragraph('- Vision Support: Inline Markdown image extraction mapped to multimodal LLMs.', style='List Bullet')
doc.add_paragraph('- Dashboard Modules: Neural Engine Console, App Widgets (Chameleon parameters), Upload RAG, Profile.', style='List Bullet')

doc.add_heading('5. Conclusion', level=1)
doc.add_paragraph('Infera Core demonstrates a state-of-the-art implementation of modern web and AI technologies. Its architecture is highly modular and prepared for heavy concurrency typical of LLM-based applications.')

doc.save(r'd:\projects\inferacore\Infera_Core_Technical_Report.docx')
print("Successfully generated Word document.")
